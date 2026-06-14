import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import get_settings


def build_docx(markdown_cv: str, job_id: int) -> str:
    """Convert markdown CV text → .docx. Returns the saved file path."""
    settings = get_settings()
    os.makedirs(settings.generated_cvs_dir, exist_ok=True)

    doc = Document()
    _set_margins(doc)
    _apply_styles(doc)

    lines = markdown_cv.strip().splitlines()
    for line in lines:
        line = line.rstrip()
        if not line:
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run(line[2:])
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=1)
            run = p.runs[0] if p.runs else p.add_run(line[3:])
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
            _add_bottom_border(p)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(11)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, line[2:])
            p.paragraph_format.left_indent = Inches(0.2)
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, line)

    path = os.path.join(settings.generated_cvs_dir, f"cv_job_{job_id}.docx")
    doc.save(path)
    return path


def _set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def _apply_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)


def _add_bottom_border(paragraph) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "162139")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_inline_formatting(paragraph, text: str) -> None:
    # Handle **bold** inline markdown
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
